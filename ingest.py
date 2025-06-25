#!/usr/bin/env python3
"""
Medical Records Summary Tool
Processes medical records and generates structured AI summaries using OpenAI's GPT-4.1-nano model
via the new Responses API.
"""

import os
import json
import argparse
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional
import logging

# External dependencies
import openai
from openai import OpenAI
import PyPDF2
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import hashlib
from io import BytesIO
import tiktoken
try:
    import fitz  # PyMuPDF for better PDF handling
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF not installed. OCR for scanned PDFs will be limited.")

# Configure Tesseract path for Windows
import platform
if platform.system() == 'Windows':
    # Try common installation paths
    import os
    possible_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\NathanEichert\AppData\Local\Tesseract-OCR\tesseract.exe',
    ]
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break
    else:
        print("Warning: Tesseract not found in common locations. Please set the path manually.")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles extraction of text from various document formats."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,  # python-docx can handle .doc files
            '.txt': self._extract_text,
            '.md': self._extract_text,
        }
        
    def extract_text(self, filepath: Path) -> str:
        """Extract text from a file based on its extension."""
        ext = filepath.suffix.lower()
        
        if ext not in self.supported_formats:
            logger.warning(f"Unsupported format: {ext} for file {filepath}")
            return ""
            
        try:
            return self.supported_formats[ext](filepath)
        except Exception as e:
            logger.error(f"Error processing {filepath}: {e}")
            return ""
    
    def _extract_pdf(self, filepath: Path) -> str:
        """
        Extract text from a PDF file page by page with clear page markers.
        Returns text with explicit page separators for accurate page citation.
        """
        text = ""
        
        # If a cached version exists (meaning we've OCR'd it before), use it.
        cache_path = filepath.with_suffix('.pdf.ocr_cache.txt')
        if cache_path.exists():
            logger.info(f"Using OCR cache for {filepath}")
            return cache_path.read_text(encoding='utf-8')

        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF not installed. Using legacy PyPDF2 for text extraction. Per-page OCR is disabled.")
            try:
                with open(filepath, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n=== PAGE {page_num + 1} START ===\n{page_text}\n=== PAGE {page_num + 1} END ===\n"
                return text
            except Exception as e:
                logger.error(f"Legacy PDF extraction error for {filepath}: {e}")
                return ""

        # --- Main logic using PyMuPDF for per-page text/OCR ---
        ocr_performed_on_any_page = False
        try:
            pdf_document = fitz.open(str(filepath))
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_number = page_num + 1
                
                # 1. Try to extract text directly from the page.
                page_text = page.get_text().strip()
                
                # 2. If text is sparse (e.g., <20 chars), it's likely a scanned image. OCR it.
                if len(page_text) < 20:
                    logger.info(f"Page {page_number} of {filepath} has little/no text; attempting OCR.")
                    ocr_performed_on_any_page = True
                    
                    try:
                        # Render page to an image for better OCR results.
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x DPI
                        img_data = pix.tobytes("png")
                        image = Image.open(BytesIO(img_data))
                        
                        ocr_text = pytesseract.image_to_string(image)
                        
                        if ocr_text.strip():
                            text += f"\n=== PAGE {page_number} START ===\n{ocr_text}\n=== PAGE {page_number} END ===\n"
                        else:
                            text += f"\n=== PAGE {page_number} START ===\n[No text found on this page]\n=== PAGE {page_number} END ===\n"
                    except Exception as ocr_error:
                        logger.error(f"OCR failed for page {page_number} of {filepath}: {ocr_error}")
                        text += f"\n=== PAGE {page_number} START ===\n[OCR failed for this page]\n=== PAGE {page_number} END ===\n"
                else:
                    # 3. If text was found, use it with clear page markers.
                    text += f"\n=== PAGE {page_number} START ===\n{page_text}\n=== PAGE {page_number} END ===\n"
                    
            pdf_document.close()
            
            # If we performed any OCR, cache the full text content for future runs.
            if ocr_performed_on_any_page:
                logger.info(f"Caching OCR results for {filepath} to {cache_path}")
                cache_path.write_text(text, encoding='utf-8')
                
        except Exception as e:
            logger.error(f"PDF extraction error for {filepath}: {e}")
        
        return text
    
    def _extract_docx(self, filepath: Path) -> str:
        """Extract text from DOCX file with page estimation."""
        try:
            doc = Document(filepath)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Estimate pages based on content (roughly 500 words per page)
            total_text = "\n".join(paragraphs)
            word_count = len(total_text.split())
            estimated_pages = max(1, (word_count // 500) + 1)
            
            if estimated_pages == 1:
                return f"\n=== PAGE 1 START ===\n{total_text}\n=== PAGE 1 END ===\n"
            else:
                # Split content across estimated pages
                lines_per_page = max(1, len(paragraphs) // estimated_pages)
                result = ""
                
                for page_num in range(estimated_pages):
                    start_idx = page_num * lines_per_page
                    end_idx = start_idx + lines_per_page if page_num < estimated_pages - 1 else len(paragraphs)
                    
                    if start_idx < len(paragraphs):
                        page_content = "\n".join(paragraphs[start_idx:end_idx])
                        result += f"\n=== PAGE {page_num + 1} START ===\n{page_content}\n=== PAGE {page_num + 1} END ===\n"
                
                return result
                
        except Exception as e:
            logger.error(f"DOCX extraction error for {filepath}: {e}")
            return ""
    
    def _extract_text(self, filepath: Path) -> str:
        """Extract text from plain text file."""
        try:
            return filepath.read_text(encoding='utf-8')
        except Exception as e:
            logger.error(f"Text extraction error for {filepath}: {e}")
            return ""
    

class MedicalRecordsSummarizer:
    """Generates AI summaries of medical records using OpenAI's Responses API."""
    
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)
        self.model = "gpt-4.1-nano-2025-04-14"  # Using the cheap, fast nano model
        # Initialize tokenizer for accurate token counting
        # GPT-4.1 uses the same tokenizer as GPT-4
        self.encoding = tiktoken.encoding_for_model("gpt-4")
    
    def _get_page_count(self, filepath: Path = None) -> int:
        """Estimate page count for a document."""
        if not filepath or not filepath.exists():
            return 1
        
        ext = filepath.suffix.lower()
        if ext == '.pdf':
            try:
                if PYMUPDF_AVAILABLE:
                    import fitz
                    doc = fitz.open(str(filepath))
                    count = len(doc)
                    doc.close()
                    return count
                else:
                    with open(filepath, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        return len(pdf_reader.pages)
            except:
                return 1
        elif ext in ['.docx', '.doc']:
            # Rough estimate for Word docs
            try:
                size_kb = filepath.stat().st_size / 1024
                return max(1, int(size_kb / 20))  # ~20KB per page estimate
            except:
                return 1
        else:
            return 1
    
    def _make_hashable(self, item):
        """
        Convert model-generated dicts/lists into a deterministic string
        so they can live in sets and JSON side-by-side.
        """
        if isinstance(item, dict):
            name = item.get("name") or item.get("party") or str(item)
            role = item.get("role") or item.get("position")
            return f"{name} ({role})" if role else name
        return str(item)
    
    def _validate_and_fix_page_citations(self, result: Dict, content: str, filepath: Path = None) -> Dict:
        """Validate and fix page citations in the result."""
        # Count actual pages in content
        page_markers = content.count("=== PAGE") // 2  # Each page has START and END
        actual_page_count = max(1, page_markers)
        
        # Update page count if it's wrong
        if result.get("page_count", 0) != actual_page_count:
            logger.info(f"Correcting page count from {result.get('page_count', 0)} to {actual_page_count}")
            result["page_count"] = actual_page_count
        
        # Validate and fix page citations in records
        if "records" in result and isinstance(result["records"], list):
            for record in result["records"]:
                if "pages_found_on" in record:
                    pages = record["pages_found_on"]
                    
                    # Ensure it's a list
                    if not isinstance(pages, list):
                        pages = [pages] if pages else [1]
                    
                    # Ensure all page numbers are valid integers within range
                    valid_pages = []
                    for page in pages:
                        try:
                            page_num = int(page)
                            if 1 <= page_num <= actual_page_count:
                                valid_pages.append(page_num)
                            else:
                                logger.warning(f"Invalid page number {page_num}, document has {actual_page_count} pages")
                        except (ValueError, TypeError):
                            logger.warning(f"Invalid page number format: {page}")
                    
                    # If no valid pages found, default to page 1
                    if not valid_pages:
                        valid_pages = [1]
                        logger.warning("No valid page numbers found, defaulting to page 1")
                    
                    record["pages_found_on"] = sorted(list(set(valid_pages)))
                else:
                    # Add default page if missing
                    record["pages_found_on"] = [1]
                    logger.info("Added missing page citation, defaulting to page 1")
        
        return result
    
    def _improve_page_citations_with_content_analysis(self, result: Dict, content: str) -> Dict:
        """Improve page citations by analyzing where quoted text actually appears in the content."""
        if "records" not in result or not isinstance(result["records"], list):
            return result
        
        # Extract page content mapping
        page_contents = {}
        lines = content.split('\n')
        current_page = None
        current_content = []
        
        for line in lines:
            if line.strip().startswith("=== PAGE") and "START ===" in line:
                # Extract page number
                try:
                    page_num = int(line.split("PAGE")[1].split("START")[0].strip())
                    current_page = page_num
                    current_content = []
                except (ValueError, IndexError):
                    continue
            elif line.strip().startswith("=== PAGE") and "END ===" in line:
                if current_page is not None:
                    page_contents[current_page] = '\n'.join(current_content)
                current_page = None
                current_content = []
            elif current_page is not None:
                current_content.append(line)
        
        # Analyze each record for better page citations
        for record in result["records"]:
            if "notes_from_visit" in record and record["notes_from_visit"]:
                quoted_text = record["notes_from_visit"].strip()
                if len(quoted_text) > 10:  # Only analyze substantial quotes
                    found_pages = []
                    
                    # Check which pages contain this quoted text
                    for page_num, page_content in page_contents.items():
                        # Look for the quote or significant portions of it
                        quote_words = quoted_text.lower().split()
                        if len(quote_words) >= 3:  # Look for 3+ word sequences
                            for i in range(len(quote_words) - 2):
                                three_word_sequence = ' '.join(quote_words[i:i+3])
                                if three_word_sequence in page_content.lower():
                                    found_pages.append(page_num)
                                    break
                    
                    if found_pages:
                        # Update page citations with found pages
                        original_pages = record.get("pages_found_on", [])
                        improved_pages = sorted(list(set(found_pages)))
                        
                        if improved_pages != original_pages:
                            logger.info(f"Improved page citation from {original_pages} to {improved_pages} based on quoted text analysis")
                            record["pages_found_on"] = improved_pages
        
        return result

    def count_tokens(self, text: str) -> int:
        """Count tokens in a text string."""
        try:
            return len(self.encoding.encode(text))
        except Exception as e:
            logger.warning(f"Token counting failed: {e}, using approximate count")
            # Fallback to rough estimate if tiktoken fails
            return len(text) // 4
        
    def summarize_document(self, filename: str, content: str, doc_type: str, filepath: Path = None, case_prompt: str = "") -> Dict[str, Any]:
        """Generate a structured summary of a medical document with per-record extraction."""
        
        # Count tokens in the content
        token_count = self.count_tokens(content)
        
        if not content or len(content.strip()) < 50:
            return {
                "filename": filename,
                "dates_covered": "Unknown",
                "hospital_facility_provider": "Unknown", 
                "page_count": self._get_page_count(filepath),
                "overall_summary": "Document too short or empty",
                "records": [],
                "token_count": token_count
            }
        
        # Truncate very long documents to fit context window
        max_chars = 50000  # Conservative limit for nano model
        is_truncated = False
        if len(content) > max_chars:
            # Find the last complete page marker before truncation
            truncated_content = content[:max_chars]
            last_page_end = truncated_content.rfind("=== PAGE")
            if last_page_end > 0:
                # Find the next "END ===" after the last page marker
                end_marker = truncated_content.find("END ===", last_page_end)
                if end_marker > 0:
                    content = truncated_content[:end_marker + 7]  # Include "END ==="
                else:
                    content = truncated_content
            else:
                content = truncated_content
            
            content += "\n\n[DOCUMENT TRUNCATED - Only pages shown above were analyzed for page citations]"
            is_truncated = True
        
        prompt = f"""Analyze this medical document and extract structured information about individual medical records/visits.

IMPORTANT: The document content includes explicit page markers in the format "=== PAGE X START ===" and "=== PAGE X END ===". Use these markers to accurately determine which page(s) contain each piece of information.

Case Context: {case_prompt}
Filename: {filename}

Content:
{content}

Provide a JSON response with the following structure:

1. "filename": The document filename
2. "dates_covered": Date range that this document covers (e.g., "01/2023 - 06/2023")
3. "hospital_facility_provider": Name of the hospital, treatment facility, or care provider
4. "page_count": Total number of pages (count the page markers in the content)
5. "overall_summary": 3-5 sentence summary of the entire document
6. "records": Array of individual medical records found in the document. For each record:
   - "record_type": One of: "primary_care_visit", "emergency_room_visit", "physical_therapy_visit", "specialist_visit", "other"
   - "date_of_visit": Date of the visit/record (MM/DD/YYYY format)
   - "notes_from_visit": Direct quotes from nurses/physicians notes (quote exactly as written)
   - "medication_treatment_changes": Any changes in medication or treatment mentioned
   - "injury_condition_description": Description of injury/condition/recovery status
   - "mentions_target_injury": Boolean - is this record relevant to the case? (See RELEVANCE CRITERIA below)
   - "pages_found_on": CRITICAL - List the exact page numbers where this record information appears based on the page markers (e.g., [1, 2])
   - "record_summary": 2-3 sentence summary of this specific record

CRITICAL INSTRUCTIONS FOR PAGE CITATIONS:
- Carefully read the page markers "=== PAGE X START ===" and "=== PAGE X END ===" in the content
- For each piece of information you extract, note which page(s) it came from based on these markers
- If a medical record spans multiple pages, include all relevant page numbers in "pages_found_on"
- If you quote text from "notes_from_visit", ensure the page numbers in "pages_found_on" accurately reflect where that text appears
- Double-check that page numbers are accurate by verifying against the page markers in the content

RELEVANCE CRITERIA FOR "mentions_target_injury" FIELD:
A medical record should be marked as relevant (mentions_target_injury: true) if it meets ANY of these criteria:

1. INCIDENT MENTION: Contains any mention or description of the specific incident at issue described in the case context above

2. INJURY/TREATMENT MENTION: Contains mention or description of:
   - The specific injuries identified in the case context
   - Treatment, therapy, or medical care for those injuries
   - Symptoms, pain, or limitations related to those injuries
   - Diagnostic tests or imaging related to those injuries

3. RELATED CONDITIONS: Contains information about:
   - Pre-existing conditions that could affect the injury claim or recovery amount
   - Post-occurring conditions that developed after the incident
   - Any medical condition that may have a positive or negative impact on liability
   - Any medical condition that may affect damages in the present case
   - Baseline health status before the incident
   - Complications or secondary conditions arising from the target injuries

Examples of relevant records:
- Visits mentioning the incident or target injuries directly
- Physical therapy for the target injury area
- Pain management related to target injuries
- Pre-existing arthritis that could affect back injury claims
- Depression developing after the incident
- Prior surgeries in the same body area
- Baseline imaging showing pre-existing conditions
- Medication changes due to the incident/injuries

Be INCLUSIVE in determining relevance - if there's any reasonable connection to the case, mark it as relevant.

Focus on extracting individual visits, treatments, or distinct medical encounters as separate records. Quote medical notes exactly as they appear in the document and cite the correct page numbers."""
        
        try:
            response = self.client.responses.create(
                model=self.model,
                input=prompt,
                temperature=0.3  # Lower temperature for factual extraction
                # Remove max_tokens - not supported in Responses API
            )
            
            # Extract the response text
            output_text = response.output_text if hasattr(response, 'output_text') else str(response)
            
            # Try to parse JSON from response
            try:
                # Find JSON content in the response
                import re
                json_match = re.search(r'\{.*\}', output_text, re.DOTALL)
                if json_match:
                    result = json.loads(json_match.group())
                else:
                    raise ValueError("No JSON found in response")
            except:
                # Fallback if JSON parsing fails
                result = {
                    "filename": filename,
                    "dates_covered": "Unknown",
                    "hospital_facility_provider": "Unknown",
                    "page_count": self._get_page_count(filepath),
                    "overall_summary": output_text[:200],
                    "records": []
                }
            
            # Ensure required fields exist
            if "filename" not in result:
                result["filename"] = filename
            if "page_count" not in result:
                result["page_count"] = self._get_page_count(filepath)
            
            result["token_count"] = token_count
            result["full_text_tokens"] = token_count  # Store original document token count
            
            # Validate and fix page citations
            result = self._validate_and_fix_page_citations(result, content, filepath)
            
            # Improve page citations by analyzing quoted text
            result = self._improve_page_citations_with_content_analysis(result, content)

            return result
            
        except Exception as e:
            logger.error(f"Error summarizing {filename}: {e}")
            return {
                "filename": filename,
                "dates_covered": "Error",
                "hospital_facility_provider": "Error", 
                "page_count": self._get_page_count(filepath),
                "overall_summary": f"Error during summarization: {str(e)}",
                "records": [],
                "token_count": token_count
            }
    
    
    def summarize_folder(self, folder_path: Path, documents: List[Dict]) -> Dict[str, Any]:
        """Generate a summary for an entire folder based on document summaries."""
        if not documents:
            return {
                "folder": folder_path.name,
                "count": 0,
                "summary": "Empty folder"
            }
        
        # Aggregate information from documents
        all_parties = set()
        all_dates = set()
        for doc in documents:
            all_parties.update(doc.get("parties", []))
            all_dates.update(doc.get("dates", []))
        
        # Create a consolidated summary
        doc_summaries = "\n".join([f"- {d['filename']}: {d.get('summary', 'No summary')}" 
                                   for d in documents[:10]])  # Limit to first 10
        
        prompt = f"""Based on these document summaries from the {folder_path.name} folder, 
provide a brief overview of what this folder contains:

{doc_summaries}

Provide a comprehensive 3-5 sentence summary of the folder's contents, key themes, and relevance to the case."""
        
        try:
            response = self.client.responses.create(
                model=self.model,
                input=prompt,
                temperature=0.3
            )
            
            summary_text = response.output_text if hasattr(response, 'output_text') else str(response)
            
            return {
                "folder": folder_path.name,
                "count": len(documents),
                "summary": summary_text.strip(),
                "key_parties": list(all_parties)[:10],  # Limit to top 10
                "key_dates": list(all_dates)[:10]
            }
            
        except Exception as e:
            logger.error(f"Error summarizing folder {folder_path}: {e}")
            return {
                "folder": folder_path.name,
                "count": len(documents),
                "summary": f"Error during folder summarization: {str(e)}"
            }

class MedicalRecordsEngine:
    """Main engine for processing medical records files."""
    
    def __init__(self, api_key: str):
        self.processor = DocumentProcessor()
        self.summarizer = MedicalRecordsSummarizer(api_key)
        
    def should_update(self, filepath: Path, summary_data: Dict) -> bool:
        """Check if a file needs to be reprocessed."""
        if filepath.name not in summary_data.get("document_registry", {}):
            return True
            
        # Check modification time
        file_mtime = filepath.stat().st_mtime
        recorded_mtime = summary_data["document_registry"][filepath.name].get("last_modified", 0)
        
        return file_mtime > recorded_mtime
    
    def process_case(self, case_path: Path, case_prompt: str = "") -> Dict[str, Any]:
        """Process an entire medical records folder and generate comprehensive summaries."""
        logger.info(f"Processing case: {case_path}")
        
        # Check for existing summary
        summary_path = case_path / "_summary.json"
        file_index_path = case_path / "_file_index.json"
        
        # Load existing data if available
        existing_summary = {}
        if summary_path.exists():
            try:
                existing_summary = json.loads(summary_path.read_text())
                logger.info("Found existing summary, will update as needed")
            except:
                logger.warning("Could not load existing summary, starting fresh")
        
        # Initialize summary structure
        summary = {
            "case_name": case_path.name,
            "last_updated": datetime.now().isoformat(),
            "folder_summaries": {},
            "document_registry": existing_summary.get("document_registry", {}),
            "key_parties": set(),
            "key_dates": set(),
            "case_type": "",
            "important_dates": {}
        }
        
        # Process all subfolders in the medical records directory
        subfolders = []
        for item in case_path.iterdir():
            if item.is_dir() and not item.name.startswith("_"):
                subfolders.append(item.name)
        
        # If no subfolders, process files in the root directory
        if not subfolders:
            subfolders = ["root"]
        
        for folder_name in subfolders:
            if folder_name == "root":
                folder_path = case_path
                logger.info(f"Processing files in root directory")
            else:
                folder_path = case_path / folder_name
                if not folder_path.exists():
                    logger.warning(f"Folder {folder_name} not found")
                    continue
                
            folder_documents = []
            
            # Process all files in the folder - focus on PDFs and text documents
            processed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md'}
            
            for filepath in folder_path.rglob("*"):
                if (filepath.is_file() and 
                    not filepath.name.startswith("_") and 
                    filepath.suffix.lower() in processed_extensions):
                    
                    # Check if we need to process this file
                    relative_path = filepath.relative_to(case_path).as_posix()
                    
                    if not self.should_update(filepath, summary):
                        logger.info(f"Skipping {filepath.name} - already processed")
                        if relative_path in summary["document_registry"]:
                            folder_documents.append(summary["document_registry"][relative_path])
                        continue
                    
                    # Extract text
                    logger.info(f"Processing: {filepath.name}")
                    text_content = self.processor.extract_text(filepath)
                    
                    if text_content:
                        # Generate summary
                        doc_summary = self.summarizer.summarize_document(
                            filepath.name,
                            text_content,
                            folder_name,
                            filepath,  # Pass filepath for image handling
                            case_prompt  # Pass case prompt for context
                        )
                        
                        # Add metadata
                        doc_summary["relative_path"] = relative_path
                        doc_summary["last_modified"] = filepath.stat().st_mtime
                        doc_summary["size_bytes"] = filepath.stat().st_size
                        
                        # Update registry
                        summary["document_registry"][relative_path] = doc_summary
                        folder_documents.append(doc_summary)
                        
                        # Aggregate key information
                        summary["key_parties"].update(doc_summary.get("parties", []))
                        summary["key_dates"].update(doc_summary.get("dates", []))
            
            # Generate folder summary
            if folder_documents:
                folder_summary = self.summarizer.summarize_folder(folder_path, folder_documents)
                summary["folder_summaries"][folder_name] = folder_summary
        
        # Convert sets to lists for JSON serialization
        summary["key_parties"] = list(summary["key_parties"])
        summary["key_dates"] = list(summary["key_dates"])
        
        # Generate overall case summary with token statistics
        summary["case_summary"] = self._generate_case_summary(summary)
        
        # Calculate total tokens across all documents
        total_tokens = sum(doc.get("token_count", 0) 
                          for doc in summary["document_registry"].values())
        summary["total_case_tokens"] = total_tokens
        summary["estimated_cost"] = f"${(total_tokens / 1_000_000) * 0.10:.4f}"  # GPT-4.1-nano pricing
        
        # Save summary
        summary_path.write_text(json.dumps(summary, indent=2))
        logger.info(f"Summary saved to {summary_path}")
        
        # Save file index for quick lookups
        file_index = {
            "case_name": case_path.name,
            "total_files": len(summary["document_registry"]),
            "total_tokens": summary.get("total_case_tokens", 0),
            "last_updated": summary["last_updated"],
            "folders": {}
        }
        
        # Calculate token counts per folder
        for folder_name in subfolders:
            folder_docs = [doc for path, doc in summary["document_registry"].items() 
                          if path.startswith(f"{folder_name}/")]
            folder_tokens = sum(doc.get("token_count", 0) for doc in folder_docs)
            file_index["folders"][folder_name] = {
                "count": len(folder_docs),
                "tokens": folder_tokens
            }
        
        file_index_path.write_text(json.dumps(file_index, indent=2))
        
        return summary
    
    def _generate_case_summary(self, summary: Dict) -> str:
        """Generate an overall case summary from all collected information."""
        total_docs = len(summary["document_registry"])
        total_tokens = summary.get("total_case_tokens", 0)
        folders_with_docs = [name for name, data in summary["folder_summaries"].items() 
                            if data.get("count", 0) > 0]
        
        # Find largest documents by token count
        largest_docs = sorted(summary["document_registry"].items(), 
                            key=lambda x: x[1].get("token_count", 0), 
                            reverse=True)[:5]
        
        overview = f"""Case: {summary['case_name']}
Total Documents: {total_docs}
Total Tokens: {total_tokens:,} (≈ {total_tokens // 1000}k)
Context Window Usage: {(total_tokens / 1_000_000) * 100:.1f}% of 1M token limit
Active Folders: {', '.join(folders_with_docs)}
Key Parties: {', '.join(summary['key_parties'][:5])}

Largest Documents by Token Count:
"""
        for path, doc in largest_docs:
            overview += f"- {doc['filename']}: {doc.get('token_count', 0):,} tokens\n"
        
        # Add folder summaries
        folder_descriptions = []
        for folder, data in summary["folder_summaries"].items():
            if data.get("summary"):
                folder_descriptions.append(f"{folder}: {data['summary']}")
        
        if folder_descriptions:
            overview += "\nFolder Contents:\n" + "\n".join(folder_descriptions[:5])
        
        return overview
    
    def split_json_file(self, json_path: Path, max_words: int = 75000, max_size_mb: int = 30) -> List[Path]:
        """Split a large JSON file into smaller chunks based on word count and file size limits."""
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Calculate file size and word count
            json_text = json.dumps(data, indent=2)
            file_size_mb = len(json_text.encode('utf-8')) / (1024 * 1024)
            word_count = len(json_text.split())
            
            logger.info(f"JSON file stats - Size: {file_size_mb:.2f}MB, Words: {word_count:,}")
            
            # Check if splitting is needed
            if file_size_mb <= max_size_mb and word_count <= max_words:
                logger.info("File is within limits, no splitting needed")
                return [json_path]
            
            # Calculate number of chunks needed
            chunks_needed_size = max(1, int(file_size_mb / max_size_mb) + 1)
            chunks_needed_words = max(1, int(word_count / max_words) + 1)
            num_chunks = max(chunks_needed_size, chunks_needed_words)
            
            logger.info(f"Splitting into {num_chunks} chunks")
            
            # Split the document registry (main content) across chunks
            documents = list(data.get("document_registry", {}).items())
            chunk_size = max(1, len(documents) // num_chunks)
            
            output_files = []
            base_name = json_path.stem
            
            for i in range(num_chunks):
                start_idx = i * chunk_size
                end_idx = start_idx + chunk_size if i < num_chunks - 1 else len(documents)
                
                if start_idx >= len(documents):
                    break
                    
                chunk_documents = dict(documents[start_idx:end_idx])
                
                # Create chunk data structure
                chunk_data = {
                    "medical_records_summary": f"{data.get('case_name', 'Medical Records')} - Part {i+1} of {num_chunks}",
                    "chunk_info": {
                        "part": i + 1,
                        "total_parts": num_chunks,
                        "documents_in_chunk": len(chunk_documents),
                        "document_range": f"{start_idx + 1}-{end_idx}"
                    },
                    "last_updated": data.get("last_updated"),
                    "case_name": data.get("case_name"),
                    "document_registry": chunk_documents,
                    "folder_summaries": data.get("folder_summaries", {}),
                    "case_summary": data.get("case_summary", "")
                }
                
                # Calculate chunk statistics
                chunk_tokens = sum(doc.get("token_count", 0) for doc in chunk_documents.values())
                chunk_data["chunk_tokens"] = chunk_tokens
                chunk_data["estimated_cost"] = f"${(chunk_tokens / 1_000_000) * 0.10:.4f}"
                
                # Write chunk file
                chunk_filename = f"{base_name}_part_{i+1:02d}_of_{num_chunks:02d}.json"
                chunk_path = json_path.parent / chunk_filename
                
                with open(chunk_path, 'w', encoding='utf-8') as f:
                    json.dump(chunk_data, f, indent=2)
                
                output_files.append(chunk_path)
                
                # Log chunk info
                chunk_text = json.dumps(chunk_data, indent=2)
                chunk_size_mb = len(chunk_text.encode('utf-8')) / (1024 * 1024)
                chunk_words = len(chunk_text.split())
                logger.info(f"Created {chunk_filename} - Size: {chunk_size_mb:.2f}MB, Words: {chunk_words:,}")
            
            # Create an index file
            index_data = {
                "original_file": json_path.name,
                "split_date": datetime.now().isoformat(),
                "total_chunks": len(output_files),
                "chunk_files": [f.name for f in output_files],
                "original_stats": {
                    "size_mb": file_size_mb,
                    "word_count": word_count,
                    "total_documents": len(documents)
                },
                "split_reason": f"Exceeded limits - Size: {file_size_mb:.1f}MB > {max_size_mb}MB or Words: {word_count:,} > {max_words:,}"
            }
            
            index_path = json_path.parent / f"{base_name}_index.json"
            with open(index_path, 'w', encoding='utf-8') as f:
                json.dump(index_data, f, indent=2)
            
            logger.info(f"Created index file: {index_path.name}")
            logger.info(f"Successfully split {json_path.name} into {len(output_files)} parts")
            
            return output_files
        
    except Exception as e:
        logger.error(f"Error splitting JSON file: {e}")
        return [json_path]

class LawyerDocumentGenerator:
    """Generates Word documents for lawyers from medical records JSON data."""
    
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)
        self.model = "gpt-4.1-nano-2025-04-14"
    
    def generate_overview_document(self, json_data: Dict, client_name: str, case_prompt: str, output_path: Path) -> Path:
        """Generate AI Medical Records Overview document."""
        logger.info("Generating Medical Records Overview document...")
        
        # Generate high-level overview using AI
        overview_summary = self._generate_overview_summary(json_data, case_prompt)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'AI Medical Records Overview - {client_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Section 1: High-Level Overview
        doc.add_heading('High-Level Medical Records Overview', level=1)
        overview_para = doc.add_paragraph(overview_summary)
        
        doc.add_paragraph()  # Spacing
        
        # Section 2: Detailed File Breakdown
        doc.add_heading('Summary of Records Relevant to Injury', level=1)
        
        # Group records by filename
        files_data = self._group_records_by_file(json_data)
        
        for filename, file_info in files_data.items():
            # File header
            file_heading = doc.add_heading(f'File: {filename}', level=2)
            
            # File summary
            doc.add_paragraph(f"Date Range: {file_info.get('dates_covered', 'Unknown')}")
            doc.add_paragraph(f"Provider/Facility: {file_info.get('hospital_facility_provider', 'Unknown')}")
            doc.add_paragraph(f"Page Count: {file_info.get('page_count', 'Unknown')}")
            
            # File overall summary
            if file_info.get('overall_summary'):
                doc.add_paragraph(f"File Summary: {file_info['overall_summary']}")
            
            # Relevant records for this file
            relevant_records = [r for r in file_info.get('records', []) 
                              if r.get('mentions_target_injury', False)]
            
            if relevant_records:
                doc.add_heading('Records Relevant to Case:', level=3)
                
                for i, record in enumerate(relevant_records, 1):
                    # Record header
                    record_para = doc.add_paragraph()
                    record_run = record_para.add_run(f"Record #{i} - {record.get('record_type', 'Unknown').replace('_', ' ').title()}")
                    record_run.bold = True
                    
                    # Record details
                    doc.add_paragraph(f"Date: {record.get('date_of_visit', 'Unknown')}")
                    doc.add_paragraph(f"Pages: {', '.join(map(str, record.get('pages_found_on', [])))}")
                    
                    if record.get('notes_from_visit'):
                        notes_para = doc.add_paragraph()
                        notes_run = notes_para.add_run("Medical Notes: ")
                        notes_run.bold = True
                        notes_para.add_run(f'"{record["notes_from_visit"]}"')
                    
                    if record.get('medication_treatment_changes'):
                        med_para = doc.add_paragraph()
                        med_run = med_para.add_run("Treatment Changes: ")
                        med_run.bold = True
                        med_para.add_run(record['medication_treatment_changes'])
                    
                    if record.get('injury_condition_description'):
                        injury_para = doc.add_paragraph()
                        injury_run = injury_para.add_run("Condition Description: ")
                        injury_run.bold = True
                        injury_para.add_run(record['injury_condition_description'])
                    
                    if record.get('record_summary'):
                        summary_para = doc.add_paragraph()
                        summary_run = summary_para.add_run("Record Summary: ")
                        summary_run.bold = True
                        summary_para.add_run(record['record_summary'])
                    
                    doc.add_paragraph()  # Spacing between records
            else:
                doc.add_paragraph("No records in this file were identified as relevant to the case.")
            
            doc.add_paragraph()  # Spacing between files
        
        # Save document
        doc_path = output_path / f"AI Medical Records Overview - {client_name}.docx"
        doc.save(str(doc_path))
        logger.info(f"Medical Records Overview saved to: {doc_path}")
        
        return doc_path
    
    def generate_chronology_document(self, json_data: Dict, client_name: str, output_path: Path) -> Path:
        """Generate AI Medical Chronology document with table."""
        logger.info("Generating Medical Chronology document...")
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'AI Medical Chronology - {client_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Collect all records from all files
        all_records = self._collect_all_records(json_data)
        
        # Sort by date
        sorted_records = self._sort_records_by_date(all_records)
        
        # Create table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Date', 'Kind of Record', 'Subject of Visit', 'Relevant to Case?', 'Summary', 'File Name', 'Page #']
        
        for i, header in enumerate(headers):
            if i < len(header_cells):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Add data rows
        for record in sorted_records:
            row_cells = table.add_row().cells
            
            # Date
            row_cells[0].text = record.get('date_of_visit', 'Unknown')
            
            # Kind of Record
            record_type = record.get('record_type', 'other').replace('_', ' ').title()
            row_cells[1].text = record_type
            
            # Subject of Visit (injury/condition description)
            subject = record.get('injury_condition_description', '')
            if not subject and record.get('notes_from_visit'):
                # Use first 100 chars of notes if no condition description
                subject = record['notes_from_visit'][:100] + '...' if len(record['notes_from_visit']) > 100 else record['notes_from_visit']
            row_cells[2].text = subject or 'Not specified'
            
            # Relevant to Case?
            relevant = 'Yes' if record.get('mentions_target_injury', False) else 'No'
            row_cells[3].text = relevant
            
            # Summary
            summary = record.get('record_summary', '')
            if not summary:
                # Create brief summary from available data
                summary_parts = []
                if record.get('medication_treatment_changes'):
                    summary_parts.append(f"Treatment: {record['medication_treatment_changes']}")
                if record.get('injury_condition_description'):
                    summary_parts.append(f"Condition: {record['injury_condition_description']}")
                summary = '; '.join(summary_parts) or 'No summary available'
            
            # Limit summary length for table
            if len(summary) > 200:
                summary = summary[:200] + '...'
            row_cells[4].text = summary
            
            # File Name
            row_cells[5].text = record.get('source_filename', 'Unknown')
            
            # Page Number
            pages = record.get('pages_found_on', [])
            row_cells[6].text = ', '.join(map(str, pages)) if pages else 'Unknown'
        
        # Add summary paragraph
        doc.add_paragraph()
        summary_para = doc.add_paragraph(f"Total Records: {len(sorted_records)}")
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        relevant_count = sum(1 for r in sorted_records if r.get('mentions_target_injury', False))
        relevant_para = doc.add_paragraph(f"Records Relevant to Case: {relevant_count}")
        relevant_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        doc_path = output_path / f"AI Med. Chronology - {client_name}.docx"
        doc.save(str(doc_path))
        logger.info(f"Medical Chronology saved to: {doc_path}")
        
        return doc_path
    
    def _generate_overview_summary(self, json_data: Dict, case_prompt: str) -> str:
        """Generate high-level overview summary using AI."""
        try:
            # Create summary of the JSON data for the prompt
            total_files = len(json_data.get('document_registry', {}))
            total_records = sum(len(doc.get('records', [])) for doc in json_data.get('document_registry', {}).values())
            relevant_records = sum(1 for doc in json_data.get('document_registry', {}).values() 
                                 for record in doc.get('records', []) 
                                 if record.get('mentions_target_injury', False))
            
            prompt = f"""Based on the medical records analysis for this personal injury case, provide a 4-5 sentence high-level overview that would be valuable for the injured person's personal injury attorney.

Case Context: {case_prompt}

Medical Records Summary:
- Total files processed: {total_files}
- Total medical records/visits: {total_records}
- Records relevant to target injury: {relevant_records}

Key findings from the records analysis:
{self._extract_key_findings(json_data)}

Please provide a concise but comprehensive overview focusing on:
1. Overall medical picture and progression
2. Key medical findings related to the injury
3. Treatment patterns and effectiveness
4. Any concerning gaps or inconsistencies
5. Critical information for legal strategy

Response should be 4-5 sentences, professional, and focused on legal relevance."""

            response = self.client.responses.create(
                model=self.model,
                input=prompt,
                temperature=0.3
            )
            
            return response.output_text if hasattr(response, 'output_text') else str(response)
            
        except Exception as e:
            logger.error(f"Error generating overview summary: {e}")
            return f"Unable to generate AI overview summary. Please review the detailed records below. Error: {str(e)}"
    
    def _extract_key_findings(self, json_data: Dict) -> str:
        """Extract key findings from the JSON data for the AI prompt."""
        findings = []
        
        for filename, doc_data in json_data.get('document_registry', {}).items():
            if doc_data.get('overall_summary'):
                findings.append(f"- {filename}: {doc_data['overall_summary']}")
        
        return '\n'.join(findings[:10])  # Limit to first 10 for prompt size
    
    def _group_records_by_file(self, json_data: Dict) -> Dict:
        """Group records by their source filename."""
        return json_data.get('document_registry', {})
    
    def _collect_all_records(self, json_data: Dict) -> List[Dict]:
        """Collect all records from all files."""
        all_records = []
        
        for filename, doc_data in json_data.get('document_registry', {}).items():
            for record in doc_data.get('records', []):
                # Add source filename to record
                record_copy = record.copy()
                record_copy['source_filename'] = filename
                all_records.append(record_copy)
        
        return all_records
    
    def _sort_records_by_date(self, records: List[Dict]) -> List[Dict]:
        """Sort records by date."""
        def parse_date(date_str):
            """Parse date string into datetime object for sorting."""
            if not date_str or date_str == 'Unknown':
                return datetime.min
            
            # Try different date formats
            formats = ['%m/%d/%Y', '%m-%d-%Y', '%Y-%m-%d', '%m/%d/%y']
            
            for fmt in formats:
                try:
                    return datetime.strptime(date_str, fmt)
                except ValueError:
                    continue
            
            # If no format works, return minimum date
            return datetime.min
        
        return sorted(records, key=lambda x: parse_date(x.get('date_of_visit', '')))
    
    def generate_lawyer_documents(self, json_data: Dict, client_name: str, case_prompt: str, output_path: Path) -> List[Path]:
        """Generate both lawyer documents."""
        generated_docs = []
        
        try:
            # Generate overview document
            overview_doc = self.generate_overview_document(json_data, client_name, case_prompt, output_path)
            generated_docs.append(overview_doc)
            
            # Generate chronology document
            chronology_doc = self.generate_chronology_document(json_data, client_name, output_path)
            generated_docs.append(chronology_doc)
            
            logger.info(f"Successfully generated {len(generated_docs)} lawyer documents")
            
        except Exception as e:
            logger.error(f"Error generating lawyer documents: {e}")
            
        return generated_docs

def main():
    parser = argparse.ArgumentParser(description="Process medical records and generate AI summaries")
    parser.add_argument("records_path", type=str, help="Path to the medical records folder")
    parser.add_argument("--api-key", type=str, help="OpenAI API key (or set OPENAI_API_KEY env var)")
    parser.add_argument("--case-prompt", type=str, default="", help="Case description and injury focus")
    parser.add_argument("--client-name", type=str, default="Client", help="Client name for document titles")
    parser.add_argument("--auto-split", action="store_true", default=True, help="Auto-split large results")
    parser.add_argument("--generate-lawyer-docs", action="store_true", default=True, help="Generate lawyer Word documents")
    parser.add_argument("--max-words", type=int, default=75000, help="Max words per file (default: 75000)")
    parser.add_argument("--max-size-mb", type=int, default=30, help="Max file size in MB (default: 30)")
    
    args = parser.parse_args()
    
    # Get API key
    api_key = args.api_key or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        logger.error("No OpenAI API key provided. Set OPENAI_API_KEY or use --api-key")
        return 1
    
    # Validate records path
    records_path = Path(args.records_path)
    if not records_path.exists():
        logger.error(f"Medical records path does not exist: {records_path}")
        return 1
    
    # Extract client name from folder if not provided
    client_name = args.client_name
    if not client_name or client_name == "Client":
        # Try to extract client name from folder path
        client_name = records_path.name
        logger.info(f"Using folder name as client name: {client_name}")
    
    # Process the medical records
    try:
        engine = MedicalRecordsEngine(api_key)
        summary = engine.process_case(records_path, args.case_prompt)
        
        logger.info("=" * 50)
        logger.info("PROCESSING COMPLETE")
        logger.info("=" * 50)
        logger.info(f"Total documents processed: {len(summary['document_registry'])}")
        logger.info(f"Total tokens: {summary.get('total_case_tokens', 0):,}")
        logger.info(f"Estimated processing cost: {summary.get('estimated_cost', 'N/A')}")
        
        # Save summary
        summary_path = records_path / '_medical_records_summary.json'
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)
        logger.info(f"Summary saved to: {summary_path}")
        
        # Generate lawyer documents if requested
        lawyer_docs = []
        if args.generate_lawyer_docs:
            logger.info("Generating lawyer documents...")
            try:
                doc_generator = LawyerDocumentGenerator(api_key)
                lawyer_docs = doc_generator.generate_lawyer_documents(
                    summary, 
                    client_name, 
                    args.case_prompt, 
                    records_path
                )
                
                if lawyer_docs:
                    logger.info("Lawyer documents generated:")
                    for doc_path in lawyer_docs:
                        logger.info(f"  - {doc_path.name}")
                else:
                    logger.warning("No lawyer documents were generated")
                    
            except Exception as e:
                logger.error(f"Error generating lawyer documents: {e}")
        
        # Auto-split if requested and needed
        if args.auto_split:
            logger.info("Checking if file splitting is needed...")
            split_files = engine.split_json_file(summary_path, args.max_words, args.max_size_mb)
            if len(split_files) > 1:
                logger.info(f"File was split into {len(split_files)} parts")
                for split_file in split_files:
                    logger.info(f"  - {split_file.name}")
        
        # Print case overview
        print("\nMEDICAL RECORDS OVERVIEW:")
        print(summary["case_summary"])
        
    except Exception as e:
        logger.error(f"Fatal error during processing: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())